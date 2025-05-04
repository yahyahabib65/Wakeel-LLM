#region import libraries
import streamlit as st
import os
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline, BitsAndBytesConfig
from peft import PeftModel

from langchain_community.llms import HuggingFacePipeline
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_community.vectorstores import Chroma
import google.generativeai as genai
from fpdf import FPDF
import datetime
from docx import Document
import base64
import streamlit.components.v1 as components


# #endregion
# def set_bg_from_local(image_file):
#     with open(image_file, "rb") as image:
#         encoded = base64.b64encode(image.read()).decode()

#     css = f"""
#     <style>
#     .stApp {{
#         background-image: url("data:image/png;base64,{encoded}");
#         background-size: cover;
#         background-position: center;
#         background-repeat: no-repeat;
#     }}
#     </style>
#     """
#     st.markdown(css, unsafe_allow_html=True)

# # Usage
# set_bg_from_local("demobackground.jpeg")

#region models and helper functions with initialization
#region lora


@st.cache_resource
def load_model():
    base_model_name = "TinyLlama/TinyLlama-1.1B-intermediate-step-1431k-3T"
    base_model = AutoModelForCausalLM.from_pretrained(base_model_name, device_map="auto", torch_dtype=torch.float16)
    tokenizer = AutoTokenizer.from_pretrained("Frontend\\tinyllama_lora_muslim_family_law")

    model = PeftModel.from_pretrained(base_model, "Frontend\\tinyllama_lora_muslim_family_law")
    model.eval()
    return model, tokenizer

model, tokenizer = load_model()

#endregion

#region RAG
# Load the RAG model and vector store
@st.cache_resource
def load_rag_model():
    # Setup Gemini
    genai.configure(api_key="AIzaSyBmpIwMg_LnNnynv6R0YW7430BJTdUX1iI")

    # Load documents
    file_path = "raw_data\\image_pdf\\family_law_manual.pdf"
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist. Please check the path.")

    loader = PyPDFLoader(file_path)
    print("Loading PDF Loader")
    documents = loader.load()
    print("Documents loaded successfully.")
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    print("Splitting documents...")
    document_chunks = text_splitter.split_documents(documents)
    print("Documents split successfully.")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    print("Embeddings loaded successfully.")
    vector_store = Chroma.from_documents(document_chunks, embeddings, persist_directory=None)
    print("Vector store loaded successfully.")
    return vector_store  # No model return now


#endregion
# Helper function to generate model response

def generate_response(prompt_text):
    inputs = tokenizer(prompt_text, return_tensors="pt").to(model.device)
    outputs = model.generate(**inputs, max_new_tokens=200)
    reply = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return reply


def generate_rag_response(prompt_text, vector_store):
    retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 2})
    relevant_docs = retriever.invoke(prompt_text)

    combined_input = (
        f"You are a legal assistant. Based on the following documents, answer the question:\n\n"
        + "\n\n".join([doc.page_content for doc in relevant_docs])
        + f"\n\nQuestion: {prompt_text}\nAnswer in simple, easy language."
    )

    model = genai.GenerativeModel('gemini-2.0-flash')
    response = model.generate_content(combined_input)

    return response.text

#endregion

#region streamlit uilibraries
# ─────────────────────────────────────────────────────────────────────────────
# 1) Include Font Awesome (for avatar/icon) – optional
st.markdown(
    '<link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.4/css/all.min.css" rel="stylesheet">',
    unsafe_allow_html=True,
)
#endregion

# ─────────────────────────────────────────────────────────────────────────────
# 2) Session-state initialization
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = {1: []}     # chat_id → list of (role, text)
if "active_chat" not in st.session_state:
    st.session_state.active_chat = 1

# ─────────────────────────────────────────────────────────────────────────────
# 3) Sidebar: select or create chats
with st.sidebar:
    st.title("WakeelGPT")
    st.header("💬 Chats")

    # List existing chats
    chat_ids = list(st.session_state.chat_sessions.keys())
    choice = st.radio(
        "Select chat",
        chat_ids,
        index=chat_ids.index(st.session_state.active_chat),
    )
    st.session_state.active_chat = choice

    st.markdown("---")
    if st.button("➕ New Chat"):
        new_id = max(chat_ids) + 1
        st.session_state.chat_sessions[new_id] = []
        st.session_state.active_chat = new_id
        st.rerun()

# ─────────────────────────────────────────────────────────────────────────────

tab1, tab2, tab3 = st.tabs(["💼 Legal Consulting", "📄 Draft Generator","Citations"])
with tab1:
    chat_container = st.container()

    # Display history
    for role, text in st.session_state.chat_sessions[st.session_state.active_chat]:
        with chat_container:
            st.chat_message(role).write(text)

    # Chat input outside the container for visibility
    
    prompt = st.chat_input(placeholder="Your message", key=f"chat_input_{st.session_state.active_chat}",accept_file=False)
    if prompt:
        # Save user message
        st.session_state.chat_sessions[st.session_state.active_chat].append(("user", prompt))

        with chat_container:
            st.chat_message("user").write(prompt)

        # Generate response
        reply = generate_response(prompt)
        # reply = "This is a placeholder response. Please implement the actual model response generation."

        # Save assistant message
        st.session_state.chat_sessions[st.session_state.active_chat].append(("assistant", reply))

        with chat_container:
            # pass
            st.chat_message("assistant", avatar=":material/gavel:").write(reply)

with tab3:

    chat_container = st.container()

    # Initialize the chat session for Tab 2
    if st.session_state.active_chat not in st.session_state.chat_sessions:
        st.session_state.chat_sessions[st.session_state.active_chat] = []

    # Display chat history
    for role, text in st.session_state.chat_sessions[st.session_state.active_chat]:
        with chat_container:
            st.chat_message(role).write(text)

    # User input field for the query
    prompt = st.chat_input(placeholder="Ask about family law in Pakistan", key=f"chat_input_t2_{st.session_state.active_chat}")
    
    if prompt:
        # Save user message
        st.session_state.chat_sessions[st.session_state.active_chat].append(("user", prompt))

        with chat_container:
            st.chat_message("user").write(prompt)

        # Load the model and vector store only for Tab 2
        vector_store = load_rag_model()

        # Generate response using the RAG function
        reply = generate_rag_response(prompt, vector_store)
        # reply = "This is a placeholder response. Please implement the actual RAG response generation."

        # Save assistant message
        st.session_state.chat_sessions[st.session_state.active_chat].append(("assistant", reply))

        with chat_container:
            st.chat_message("assistant", avatar=":material/gavel:").write(reply)

# with tab2:
#     # Container for messages
#     chat_container = st.container()

#     # Display history
#     for role, text in st.session_state.chat_sessions[st.session_state.active_chat]:
#         with chat_container.chat_message(role):
#             chat_container.write(text)

#     # with st.sidebar:
#     messages = st.container()
#     if prompt := st.chat_input(placeholder="Your message",accept_file=True,key="citations"):
#         messages.chat_message("user").write(prompt.text)
#         # messages.chat_message("assistant",avatar=":material/gavel:").write(f"Echo: {prompt.text}")

# with tab2:
#     st.header("📄 Draft Generator")
#     st.write("Answer the questions step-by-step to generate a legal document.")

#     # Initialize session state for draft generation
#     if "draft_conversation" not in st.session_state:
#         st.session_state.draft_conversation = []  # Stores the conversation history
#     if "draft_complete" not in st.session_state:
#         st.session_state.draft_complete = False

#     # Display conversation history
#     for role, text in st.session_state.draft_conversation:
#         with st.chat_message(role):
#             st.write(text)

#     # If the draft is not complete, continue the conversation
#     if not st.session_state.draft_complete:
#         # User input for the current question
#         user_input = st.chat_input(placeholder="Your answer...")

#         if user_input:
#             # Save user input
#             st.session_state.draft_conversation.append(("user", user_input))

#             # Generate the next question or the draft
#             with st.spinner("Processing..."):
#                 try:
#                     # Combine the conversation history into a single prompt
#                     conversation_history = "\n".join(
#                         [f"{role.capitalize()}: {text}" for role, text in st.session_state.draft_conversation]
#                     )
#                     prompt = (
#                         # f"You are a legal assistant helping to draft a legal document. "
#                         # f"Based on the following conversation, ask the next relevant question or generate the draft:\n\n"
#                         # f"{conversation_history}\n\n"
#                         # f"Assistant:"
#                         "You are a legal assistant drafting a Khula legal document under Pakistani Muslim Family Law.\nYour task is to ask the user questions one by one to gather all necessary legal information.\nAsk only one specific question at a time.\n\nSo far, this is the conversation:"+ conversation_history + "\n\nNow ask the next relevant question."
#                     )

#                     # Use the LLM to generate the next response
#                     response = generate_response(prompt)

#                     # Check if the response is a draft or a question
#                     if "IN THE FAMILY COURT" in response or "PETITION" in response:
#                         st.session_state.draft_complete = True
#                         st.session_state.draft_conversation.append(("assistant", response))
#                     else:
#                         st.session_state.draft_conversation.append(("assistant", response))
#                 except Exception as e:
#                     st.error(f"Error: {e}")

#     # If the draft is complete, display it and provide download options
#     if st.session_state.draft_complete:
#         draft = st.session_state.draft_conversation[-1][1]  # The last assistant message is the draft
#         st.subheader("Generated Draft")
#         st.code(draft, language="text")

#         # Provide download options
#         def generate_pdf(text, filename="Legal_Draft.pdf"):
#             pdf = FPDF()
#             pdf.set_auto_page_break(auto=True, margin=15)
#             pdf.add_page()
#             pdf.set_font("Arial", size=12)
#             pdf.multi_cell(0, 10, text)
#             pdf.output(filename)
#             return filename

#         def generate_docx(text, filename="Legal_Draft.docx"):
#             doc = Document()
#             doc.add_paragraph(text)
#             doc.save(filename)
#             return filename

#         def generate_txt(text, filename="Legal_Draft.txt"):
#             with open(filename, "w") as file:
#                 file.write(text)
#             return filename

#         # Generate files
#         pdf_filename = generate_pdf(draft)
#         docx_filename = generate_docx(draft)
#         txt_filename = generate_txt(draft)

#         # Provide download buttons
#         with open(pdf_filename, "rb") as f:
#             st.download_button("⬇️ Download PDF", f, file_name=pdf_filename)

#         with open(docx_filename, "rb") as f:
#             st.download_button("⬇️ Download DOCX", f, file_name=docx_filename)

#         with open(txt_filename, "rb") as f:
#             st.download_button("⬇️ Download TXT", f, file_name=txt_filename)

#         # Reset option
#         if st.button("Start Over"):
#             st.session_state.draft_conversation = []
#             st.session_state.draft_complete = False
#             st.rerun()


# with tab2:
#     # Initialize the draft session for Tab 2
#     if "draft_step" not in st.session_state:
#         st.session_state.draft_step = 0
#         st.session_state.draft_answers = {}
#         st.session_state.draft_finished = False
#         st.session_state.chat_sessions[st.session_state.active_chat] = []

#     # Draft type (e.g., Khula Petition)
#     draft_options = [
#         ("wife_name", "👩 What is the wife's full name?"),
#         ("husband_name", "🧔 What is the husband's full name?"),
#         ("marriage_date", "📅 When did the marriage take place?"),
#         ("place_of_marriage", "📍 Where was the marriage held?"),
#         ("reason_for_khula", "💔 What is the reason for seeking Khula?"),
#         ("mehr_details", "💰 What were the Mehr details?"),
#         ("children_details", "👶 Are there any children?"),
#         ("prayer", "🙏 What relief is being sought?")
#     ]
    
#     # Display chat history for the draft generation process
#     for role, text in st.session_state.chat_sessions[st.session_state.active_chat]:
#         with chat_container:
#             st.chat_message(role).write(text)

#     # Show the next question based on the draft step
#     if not st.session_state.draft_finished:
#         current_field, question = draft_options[st.session_state.draft_step]
#         with st.chat_message("assistant"):
#             st.markdown(question)

#         user_input = st.chat_input("Your answer...")

#         if user_input:
#             st.session_state.draft_answers[current_field] = user_input
#             st.session_state.chat_sessions[st.session_state.active_chat].append(("user", user_input))
#             st.session_state.chat_sessions[st.session_state.active_chat].append(("assistant", question))

#             # Move to the next step
#             st.session_state.draft_step += 1

#             # If all questions have been answered, generate the draft
#             if st.session_state.draft_step >= len(draft_options):
#                 st.session_state.draft_finished = True
#                 draft = generate_draft("Khula Petition", st.session_state.draft_answers)

#                 # Save and show the draft
#                 with open(f"Khula_Petition_{datetime.date.today()}.docx", "w") as file:
#                     file.write(draft)
                
#                 st.success("✅ Your draft is ready!")
#                 st.code(draft)
#                 with open(f"Khula_Petition_{datetime.date.today()}.docx", "rb") as f:
#                     st.download_button("⬇️ Download Draft", f, file_name=f"Khula_Petition_{datetime.date.today()}.docx")

#     # If finished, ask if they want to start a new draft
#     if st.session_state.draft_finished:
#         with st.chat_message("assistant"):
#             st.markdown("Would you like to create another draft?")
#         if st.button("Create New Draft"):
#             st.session_state.draft_step = 0
#             st.session_state.draft_answers = {}
#             st.session_state.draft_finished = False
#             st.session_state.chat_sessions[st.session_state.active_chat] = []
#             st.session_state.active_chat += 1
#             st.rerun()

# with tab2:
#     # Initialize the draft session for Tab 2
#     if "draft_type" not in st.session_state:
#         st.session_state.draft_type = None
#         st.session_state.step = 0
#         st.session_state.answers = {}
#         st.session_state.finished = False
#         st.session_state.chat_history = []

#     # Draft Type Selector (only first time)
#     if not st.session_state.draft_type:
#         st.session_state.draft_type = st.selectbox("📑 Select the type of legal draft you want to generate:", ["Khula Petition", "Property Transfer After Death", "Will Deed", "Marriage Registration"])

#     # Define draft questions
#     draft_options = {
#         "Khula Petition": [
#             ("wife_name", "👩 What is the wife's full name?"),
#             ("husband_name", "🧔 What is the husband's full name?"),
#             ("marriage_date", "📅 When did the marriage take place?"),
#             ("place_of_marriage", "📍 Where was the marriage held?"),
#             ("reason_for_khula", "💔 What is the reason for seeking Khula?"),
#             ("mehr_details", "💰 What were the Mehr details?"),
#             ("children_details", "👶 Are there any children?"),
#             ("prayer", "🙏 What relief is being sought?")
#         ],
#         "Property Transfer After Death": [
#             ("deceased_name", "🪦 What is the full name of the deceased?"),
#             ("date_of_death", "📅 When did the person pass away?"),
#             ("property_details", "🏠 Describe the property to be transferred."),
#             ("legal_heirs", "👪 List the legal heirs and their shares."),
#             ("transfer_reason", "⚖️ State the reason for the transfer."),
#         ],
#         "Will Deed": [
#             ("testator_name", "👤 Full name of the person making the will:"),
#             ("testator_address", "🏠 Address of the testator:"),
#             ("property_to_be_distributed", "🏘️ Describe the property to be included in the will:"),
#             ("beneficiaries", "👨‍👩‍👧‍👦 List the names and shares of beneficiaries:"),
#             ("executor_name", "📜 Who will execute this will?")
#         ],
#         "Marriage Registration": [
#             ("bride_name", "👰 Full name of the bride:"),
#             ("groom_name", "🤵 Full name of the groom:"),
#             ("marriage_date", "📅 Date of marriage:"),
#             ("place_of_marriage", "📍 Place of marriage:"),
#             ("witnesses", "🧾 Names of two witnesses:"),
#             ("nikah_registrar", "📋 Name of the Nikah Registrar:")
#         ]
#     }

#     # Define a function to generate drafts based on the collected answers
#     def generate_draft(draft_type, data):
#         if draft_type == "Khula Petition":
#             return f"""IN THE FAMILY COURT AT [City Name]

# In the matter of:
# {data['wife_name']} (Petitioner)
# Versus
# {data['husband_name']} (Respondent)

# PETITION FOR KHULA

# Respectfully Sheweth:

# 1. That the petitioner was married to the respondent on {data['marriage_date']} at {data['place_of_marriage']}.
# 2. That the relationship has irretrievably broken down due to: {data['reason_for_khula']}.
# 3. That the petitioner is willing to return the mehr amount of {data['mehr_details']}.
# 4. {data['children_details']}

# PRAYER:
# {data['prayer']}

# Petitioner: {data['wife_name']}
# Date: {datetime.date.today().strftime('%d %B %Y')}
# """
#         elif draft_type == "Property Transfer After Death":
#             return f"""PROPERTY TRANSFER DECLARATION

# This is to certify that {data['deceased_name']} passed away on {data['date_of_death']}.

# The following property is to be transferred:
# {data['property_details']}

# Legal heirs entitled to the property are:
# {data['legal_heirs']}

# Reason for transfer:
# {data['transfer_reason']}

# Date: {datetime.date.today().strftime('%d %B %Y')}
# """
#         elif draft_type == "Will Deed":
#             return f"""WILL DEED

# I, {data['testator_name']}, residing at {data['testator_address']}, being of sound mind and disposing memory, do hereby declare this to be my last will and testament.

# The following property shall be distributed:
# {data['property_to_be_distributed']}

# Beneficiaries and their shares:
# {data['beneficiaries']}

# Executor of this Will:
# {data['executor_name']}

# Executed on this day: {datetime.date.today().strftime('%d %B %Y')}
# """
#         elif draft_type == "Marriage Registration":
#             return f"""APPLICATION FOR MARRIAGE REGISTRATION

# Bride: {data['bride_name']}
# Groom: {data['groom_name']}
# Date of Marriage: {data['marriage_date']}
# Place of Marriage: {data['place_of_marriage']}

# Witnesses:
# {data['witnesses']}

# Nikah Registrar: {data['nikah_registrar']}

# Submitted on: {datetime.date.today().strftime('%d %B %Y')}
# """
#         else:
#             return "❌ Draft type not recognized."

#     # Export to DOCX function
#     def export_to_docx(text, filename="Legal_Draft.docx"):
#         doc = Document()
#         doc.add_paragraph(text)
#         doc.save(filename)
#         return filename

#     # Draft Process: Ask Questions, Record Answers
#     if not st.session_state.finished:
#         questions = draft_options[st.session_state.draft_type]
#         current_field, question = questions[st.session_state.step]
        
#         with st.chat_message("assistant"):
#             st.markdown(question)

#         user_input = st.chat_input("Your answer...")

#         if user_input:
#             st.session_state.answers[current_field] = user_input
#             st.session_state.chat_history.append(("user", user_input))
#             st.session_state.chat_history.append(("assistant", question))

#             st.session_state.step += 1

#             if st.session_state.step >= len(questions):
#                 st.session_state.finished = True
#                 draft = generate_draft(st.session_state.draft_type, st.session_state.answers)
#                 filename = export_to_docx(draft, f"{st.session_state.draft_type.replace(' ', '_')}.docx")

#                 with st.chat_message("assistant"):
#                     st.success("✅ Your draft is ready!")
#                     st.code(draft)
#                     with open(filename, "rb") as f:
#                         st.download_button("⬇️ Download Draft", f, file_name=filename)

#     # Allow starting a new draft if finished
#     if st.session_state.finished:
#         with st.chat_message("assistant"):
#             st.markdown("Would you like to create another draft?")
#         if st.button("Create New Draft"):
#             st.session_state.draft_type = None
#             st.session_state.step = 0
#             st.session_state.answers = {}
#             st.session_state.finished = False
#             st.session_state.chat_history = []
#             st.session_state.draft_type = st.selectbox("📑 Select the type of legal draft you want to generate:", ["Khula Petition", "Property Transfer After Death", "Will Deed", "Marriage Registration"])
#             st.rerun()



with tab2:
    def export_to_docx(text, filename="Legal_Draft.docx"):
        doc = Document()
        doc.add_paragraph(text)
        doc.save(filename)
        return filename


    def generate_draft(draft_type, data):
        if draft_type == "Khula Petition":
            return f"""IN THE FAMILY COURT AT [City Name]

    In the matter of:
    {data['wife_name']} (Petitioner)
    Versus
    {data['husband_name']} (Respondent)

    PETITION FOR KHULA

    Respectfully Sheweth:

    1. That the petitioner was married to the respondent on {data['marriage_date']} at {data['place_of_marriage']}.
    2. That the relationship has irretrievably broken down due to: {data['reason_for_khula']}.
    3. That the petitioner is willing to return the mehr amount of {data['mehr_details']}.
    4. {data['children_details']}

    PRAYER:
    {data['prayer']}

    Petitioner: {data['wife_name']}
    Date: {datetime.date.today().strftime('%d %B %Y')}
    """
        elif draft_type == "Property Transfer After Death":
            return f"""PROPERTY TRANSFER DECLARATION

    This is to certify that {data['deceased_name']} passed away on {data['date_of_death']}.

    The following property is to be transferred:
    {data['property_details']}

    Legal heirs entitled to the property are:
    {data['legal_heirs']}

    Reason for transfer:
    {data['transfer_reason']}

    Date: {datetime.date.today().strftime('%d %B %Y')}
    """
        elif draft_type == "Will Deed":
            return f"""WILL DEED

    I, {data['testator_name']}, residing at {data['testator_address']}, being of sound mind and disposing memory, do hereby declare this to be my last will and testament.

    The following property shall be distributed:
    {data['property_to_be_distributed']}

    Beneficiaries and their shares:
    {data['beneficiaries']}

    Executor of this Will:
    {data['executor_name']}

    Executed on this day: {datetime.date.today().strftime('%d %B %Y')}
    """
        elif draft_type == "Marriage Registration":
            return f"""APPLICATION FOR MARRIAGE REGISTRATION

    Bride: {data['bride_name']}
    Groom: {data['groom_name']}
    Date of Marriage: {data['marriage_date']}
    Place of Marriage: {data['place_of_marriage']}

    Witnesses:
    {data['witnesses']}

    Nikah Registrar: {data['nikah_registrar']}

    Submitted on: {datetime.date.today().strftime('%d %B %Y')}
    """
        else:
            return "❌ Draft type not recognized."


    # 1) Initialize session state
    if "draft_type" not in st.session_state:
        st.session_state.draft_type = None
        st.session_state.step = 0
        st.session_state.answers = {}
        st.session_state.finished = False
        st.session_state.chat_history = []

    # 2) Select draft type once
    if st.session_state.draft_type is None:
        st.session_state.draft_type = st.selectbox(
            "📑 Select the type of legal draft:",
            ["Khula Petition", "Property Transfer After Death", "Will Deed", "Marriage Registration"],
        )

    # 3) Question bank
    draft_options = {
      "Khula Petition": [
         ("wife_name",      "👩 What is the wife's full name?"),
         ("husband_name",   "🧔 What is the husband's full name?"),
         ("marriage_date", "📅 When did the marriage take place?"),
        ("place_of_marriage", "📍 Where was the marriage held?"),
        ("reason_for_khula", "💔 What is the reason for seeking Khula?"),
        ("mehr_details", "💰 What were the Mehr details?"),
        ("children_details", "👶 Are there any children?"),
        ("prayer", "🙏 What relief is being sought?")
      ],
      # other types …
    }

    # 4) Single container for chat history + input
    chat_area = st.container()

    # 5) Replay history
    for role, msg in st.session_state.chat_history:
        with chat_area:
            st.chat_message(role).markdown(msg)

    # 6) If not finished, ask next question
    if not st.session_state.finished:
        field, question = draft_options[st.session_state.draft_type][st.session_state.step]
        with chat_area:
            st.chat_message("assistant").markdown(question)

        # *** Give the input widget a stable key tied to the step ***
        user_input = chat_area.chat_input(
            "Your answer…",
            key=f"answer_{st.session_state.step}"
        )

        if user_input:
            # record
            st.session_state.answers[field] = user_input
            st.session_state.chat_history.append(("user", user_input))
            st.session_state.chat_history.append(("assistant", question))
            st.session_state.step += 1

            # check completion
            if st.session_state.step >= len(draft_options[st.session_state.draft_type]):
                st.session_state.finished = True

    # 7) Generate & show draft once complete
    if st.session_state.finished:
        draft = generate_draft(st.session_state.draft_type, st.session_state.answers)
        filename = export_to_docx(draft, f"{st.session_state.draft_type}.docx")

        with chat_area:
            st.success("✅ Your draft is ready!")
            st.code(draft)
            with open(filename, "rb") as f:
                st.download_button("⬇️ Download Draft", f, file_name=filename)

        # reset button
        if st.button("🔁 Create New Draft"):
            for k in ["draft_type","step","answers","finished","chat_history"]:
                del st.session_state[k]
            st.experimental_rerun()



